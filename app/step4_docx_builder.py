import os
import json
import re
from docx import Document
from docx.shared import Pt, RGBColor

def clean_tag(text: str) -> str:
    if not isinstance(text, str):
        return ""
    text = re.sub(r'^[🔴⚫🔵]\s*', '', text)
    text = re.sub(r'^(T:|Q:|S:|A:)\s*', '', text)
    return text.strip()

def safe_get_dict(obj):
    if isinstance(obj, dict):
        return obj
    if isinstance(obj, str):
        try:
            res = json.loads(obj)
            if isinstance(res, dict):
                return res
        except Exception:
            pass
    return {}

def build_docx_from_json(json_path: str, output_docx_path: str):
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    if isinstance(data, str):
        try:
            data = json.loads(data)
        except Exception:
            data = {}

    data = safe_get_dict(data)

    doc = Document()

    COLOR_BLACK = RGBColor(0, 0, 0)
    COLOR_DARK_RED = RGBColor(181, 22, 0)
    COLOR_DARK_BLUE = RGBColor(0, 76, 128)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Helvetica Neue'
    font.size = Pt(15)

    # 1. Title
    title_p = doc.add_paragraph()
    r = title_p.add_run("ESL Picture Book Lesson Plan")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = COLOR_BLACK

    # Meta Info
    p_meta = doc.add_paragraph()
    r_meta = p_meta.add_run(
        f"Book: {data.get('book_title', '')}\n"
        f"Level: {data.get('level', 'Beginner / Elementary')}\n"
        f"Duration: {data.get('duration', '25 minutes')}\n"
    )
    r_meta.font.color.rgb = COLOR_BLACK

    # 2. Teaching Intention Analysis
    p_ti_head = doc.add_paragraph()
    r = p_ti_head.add_run("Teaching Intention Analysis")
    r.bold = True
    r.font.color.rgb = COLOR_BLACK

    ti = safe_get_dict(data.get("teaching_intention", {}))
    p_sum = doc.add_paragraph()
    p_sum.add_run("Story Summary: ").bold = True
    p_sum.add_run(f"{ti.get('story_summary', '')}\n").font.color.rgb = COLOR_BLACK

    p_goals = doc.add_paragraph()
    p_goals.add_run("Teaching Goals:\n").bold = True
    for goal_key, goal_label in [("language_goal", "Language Goal"), ("vocabulary_goal", "Vocabulary Goal"), ("thinking_goal", "Thinking Goal")]:
        rg = p_goals.add_run(f"• {goal_label}: {ti.get(goal_key, '')}\n")
        rg.font.color.rgb = COLOR_DARK_BLUE

    # 3. Part 1: Greeting & Warm-up
    p_part1 = doc.add_paragraph()
    r = p_part1.add_run("Part 1: Greeting & Warm-up (0:00–3:00) — Cover")
    r.bold = True
    r.font.color.rgb = COLOR_BLACK

    cover = safe_get_dict(data.get("cover", {}))
    questions = cover.get("questions", [])
    if isinstance(questions, list):
        for q_obj in questions:
            q_obj = safe_get_dict(q_obj)
            p_pair = doc.add_paragraph()
            rq = p_pair.add_run(f"{clean_tag(q_obj.get('question', ''))}\n")
            rq.font.color.rgb = COLOR_DARK_RED
            ra = p_pair.add_run(f"{clean_tag(q_obj.get('answer', ''))}")
            ra.font.color.rgb = COLOR_BLACK

    vocab = cover.get("vocabulary", [])
    if isinstance(vocab, list):
        for v_obj in vocab:
            v_obj = safe_get_dict(v_obj)
            p_v = doc.add_paragraph()
            rv = p_v.add_run(f"{clean_tag(v_obj.get('word', ''))} — {clean_tag(v_obj.get('meaning', ''))}")
            rv.font.color.rgb = COLOR_DARK_BLUE

    # 4. Part 2: Reading
    p_part2 = doc.add_paragraph()
    r = p_part2.add_run("Part 2: Reading (3:00–20:00)")
    r.bold = True
    r.font.color.rgb = COLOR_BLACK

    pages = data.get("pages", [])
    if isinstance(pages, list):
        for page in pages:
            if not isinstance(page, dict):
                continue
            p_num = page.get("page_number", "")
            p_quote = page.get("story_quote", "")
            
            p_head = doc.add_paragraph()
            rh = p_head.add_run(f"Page {p_num} — \"{p_quote}\"")
            rh.bold = True
            rh.font.color.rgb = COLOR_BLACK

            p_qs = page.get("questions", [])
            if isinstance(p_qs, list):
                for q_obj in p_qs:
                    q_obj = safe_get_dict(q_obj)
                    p_pair = doc.add_paragraph()
                    rq = p_pair.add_run(f"{clean_tag(q_obj.get('question', ''))}\n")
                    rq.font.color.rgb = COLOR_DARK_RED
                    ra = p_pair.add_run(f"{clean_tag(q_obj.get('answer', ''))}")
                    ra.font.color.rgb = COLOR_BLACK

            p_voc = page.get("vocabulary", [])
            if isinstance(p_voc, list):
                for v_obj in p_voc:
                    v_obj = safe_get_dict(v_obj)
                    p_v = doc.add_paragraph()
                    rv = p_v.add_run(f"{clean_tag(v_obj.get('word', ''))} — {clean_tag(v_obj.get('meaning', ''))}")
                    rv.font.color.rgb = COLOR_DARK_BLUE

    # 5. Grammar Focus
    p_gf_head = doc.add_paragraph()
    r = p_gf_head.add_run("Grammar Focus")
    r.bold = True
    r.font.color.rgb = COLOR_BLACK

    gf = safe_get_dict(data.get("grammar_focus", {}))
    p_g = doc.add_paragraph()
    
    g_point = clean_tag(gf.get('grammar_point', 'Simple Present / Past Tense Practice'))
    rg1 = p_g.add_run(f"Grammar Point: {g_point}\n")
    rg1.font.color.rgb = COLOR_DARK_BLUE
    
    g_ex = gf.get('story_example', '')
    if g_ex:
        p_g.add_run(f"Story Example: \"{g_ex}\"\n").font.color.rgb = COLOR_BLACK
        
    g_frame = clean_tag(gf.get('student_sentence_frame', ''))
    if g_frame:
        rg2 = p_g.add_run(f"Student Sentence Frame: {g_frame}\n")
        rg2.font.color.rgb = COLOR_DARK_BLUE
        
    if gf.get("other_matching_words_in_story"):
        p_g.add_run(f"Other verbs/words in story: {gf.get('other_matching_words_in_story')}\n").italic = True

    # 6. Part 3: Game & Wrap-up
    p_part3 = doc.add_paragraph()
    r = p_part3.add_run("Part 3: Game & Wrap-up (20:00–25:00)")
    r.bold = True
    r.font.color.rgb = COLOR_BLACK

    games = data.get("games", [])
    if isinstance(games, list):
        for idx, game in enumerate(games, 1):
            if not isinstance(game, dict):
                continue
            p_gh = doc.add_paragraph()
            g_name = game.get('game_type', game.get('game_name', f'QKids Interactive Game {idx}'))
            rgh = p_gh.add_run(f"Game {idx}: {g_name}")
            rgh.bold = True
            rgh.font.color.rgb = COLOR_BLACK

            p_gm = doc.add_paragraph()
            if game.get("purpose"):
                p_gm.add_run(f"Purpose: {game.get('purpose')}\n")
            p_gm.add_run(f"Target language: {game.get('target_language', '')}\n").bold = True
            
            setup = game.get("screen_setup", [])
            if isinstance(setup, list):
                p_gm.add_run("Screen setup:\n")
                for step in setup:
                    p_gm.add_run(f"  • {step}\n")
            elif isinstance(setup, str):
                p_gm.add_run(f"Screen setup: {setup}\n")

            if game.get("teacher_says"):
                p_gm.add_run(f"Teacher says: \"{game.get('teacher_says')}\"\n")
            elif game.get("teacher_action"):
                p_gm.add_run(f"Teacher action: {game.get('teacher_action')}\n")

            p_gm.add_run(f"Student output: \"{game.get('student_output', '')}\"\n").italic = True

    # 7. Story Retelling
    p_ret_head = doc.add_paragraph()
    r = p_ret_head.add_run("Story Retelling")
    r.bold = True
    r.font.color.rgb = COLOR_BLACK

    retelling = data.get("story_retelling", [])
    if isinstance(retelling, list):
        for item in retelling:
            frame_text = item.get("frame", "") if isinstance(item, dict) else str(item)
            p_r = doc.add_paragraph()
            rr = p_r.add_run(clean_tag(frame_text))
            rr.font.color.rgb = COLOR_DARK_BLUE

    os.makedirs(os.path.dirname(output_docx_path), exist_ok=True)
    doc.save(output_docx_path)
    print(f"SUCCESS! Fully populated DOCX saved to: '{output_docx_path}'")


if __name__ == "__main__":
    json_input = "test_lesson_output.json"
    output_docx = "output_lessons/Things_He_Chews_Lesson_Plan.docx"
    if os.path.exists(json_input):
        build_docx_from_json(json_input, output_docx)
    else:
        print(f"ERROR: Could not find '{json_input}'.")