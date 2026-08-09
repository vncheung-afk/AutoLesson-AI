from docx import Document
from docx.shared import RGBColor, Pt, Inches
import os
import json


BASE_DIR = os.path.dirname(
    os.path.dirname(os.path.abspath(__file__))
)

OUTPUT_FOLDER = os.path.join(
    BASE_DIR,
    "output"
)

DATA_FILE = os.path.join(
    BASE_DIR,
    "tests",
    "lesson_data.json"
)

FONT_NAME = "Arial"


RED = (255, 0, 0)
BLACK = (0, 0, 0)
BLUE = (0, 0, 255)


def setup_document(doc):

    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(15)

    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)


def format_run(run, color, bold=False, size=15):

    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*color)


def add_text(doc, text, color=BLACK, bold=False):

    p = doc.add_paragraph()

    run = p.add_run(text)

    format_run(
        run,
        color,
        bold
    )

    return p


def add_heading(doc, text, level=2):

    h = doc.add_heading(
        text,
        level=level
    )

    for run in h.runs:
        format_run(
            run,
            BLACK,
            bold=True,
            size=16 if level == 1 else 15
        )


def create_lesson_doc(data):

    os.makedirs(
        OUTPUT_FOLDER,
        exist_ok=True
    )

    filename = os.path.join(
        OUTPUT_FOLDER,
        f"{data['book_info']['title']}_Lesson.docx"
    )

    doc = Document()

    setup_document(doc)


    # Title
    add_heading(
        doc,
        "ESL Picture Book Lesson Plan",
        level=1
    )

    add_text(
        doc,
        f"Book: {data['book_info']['title']}"
    )

    add_text(
        doc,
        f"Level: {data['book_info']['level']}"
    )

    add_text(
        doc,
        f"Duration: {data['book_info']['duration']}"
    )


    # Teaching Analysis
    add_heading(
        doc,
        "Teaching Intention Analysis"
    )

    add_text(
        doc,
        "Story Summary: " + data["teaching_analysis"]["story_summary"]
    )

    add_text(
        doc,
        "Language Goal: " +
        ", ".join(data["teaching_analysis"]["language_goal"])
    )

    add_text(
        doc,
        "Vocabulary Goal: " +
        ", ".join(data["teaching_analysis"]["vocabulary_goal"])
    )

    add_text(
        doc,
        "Thinking Goal: " +
        ", ".join(data["teaching_analysis"]["thinking_goal"])
    )


    # Cover
    add_heading(
        doc,
        "Cover Discussion"
    )

    for q in data["cover"]["questions"]:

        add_text(
            doc,
            "Q: " + q["question"],
            RED
        )

        add_text(
            doc,
            "A: " + q["answer"],
            BLACK
        )


    # Pages
    add_heading(
        doc,
        "Reading"
    )

    for page in data["pages"]:

        add_heading(
            doc,
            f"Page {page['page_number']}"
        )

        if page["story_text"]:
            add_text(
                doc,
                page["story_text"]
            )


        for q in page["picture_observation"]:

            add_text(
                doc,
                "Q: " + q["question"],
                RED
            )

            add_text(
                doc,
                "A: " + q["answer"]
            )


        for q in page["comprehension_questions"]:

            add_text(
                doc,
                "Q: " + q["question"],
                RED
            )

            add_text(
                doc,
                "A: " + q["answer"]
            )


        for v in page["vocabulary"]:

            add_text(
                doc,
                f"{v['word']} - {v['meaning']}",
                BLUE
            )

            if v["example_sentence"]:
                add_text(
                    doc,
                    v["example_sentence"],
                    BLUE
                )


    # Grammar
    add_heading(
        doc,
        "Grammar Focus"
    )

    add_text(
        doc,
        data["grammar_focus"]["target"],
        BLUE
    )

    add_text(
        doc,
        data["grammar_focus"]["explanation"]
    )


    # Games
    add_heading(
        doc,
        "Games"
    )

    for game in data["games"]:

        add_text(
            doc,
            game["name"],
            BLUE,
            True
        )

        add_text(
            doc,
            "How to play: " + game["instructions"]
        )

        add_text(
            doc,
            "Student output: " + game["student_output"]
        )


    # Retelling
    add_heading(
        doc,
        "Story Retelling"
    )

    add_text(
        doc,
        data["retelling"]["structure"],
        BLUE
    )

    add_text(
        doc,
        data["retelling"]["model_answer"]
    )


    doc.save(filename)

    print("Created:", filename)


if __name__ == "__main__":

    with open(
        DATA_FILE,
        "r",
        encoding="utf-8"
    ) as f:

        lesson = json.load(f)


    create_lesson_doc(lesson)